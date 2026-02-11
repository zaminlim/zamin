# ==============================
# 설치 필요 패키지
# ==============================
# pip install yfinance matplotlib pandas numpy python-docx langchain-openai

# ==============================
# 라이브러리
# ==============================
import os
from pathlib import Path
import getpass
import yfinance as yf
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from langchain_openai import ChatOpenAI
from io import BytesIO

# ==============================
# 한글 폰트 설정
# ==============================
plt.rcParams["font.family"] = "Malgun Gothic"
plt.rcParams["axes.unicode_minus"] = False

# ==============================
# OpenAI API 키 입력
# ==============================
api_key = getpass.getpass("OpenAI API Key를 입력하세요 (입력해도 화면에는 표시되지 않습니다): ")
if not api_key:
    raise ValueError("OpenAI API Key가 입력되지 않았습니다.")

# ==============================
# GPT 초기화
# ==============================
llm = ChatOpenAI(model="gpt-4.1-mini", temperature=0, openai_api_key=api_key)

# ==============================
# 재무 데이터 함수
# ==============================
def get_annual_financials(ticker):
    stock = yf.Ticker(ticker)
    df = stock.financials.loc[["Total Revenue", "Operating Income"]].T
    df = df.sort_index().tail(3)
    df.index = df.index.year
    return df

def get_quarterly_actuals(ticker):
    stock = yf.Ticker(ticker)
    df = stock.quarterly_financials.loc[["Total Revenue", "Operating Income"]].T
    df = df.sort_index().tail(4)
    return df

def forecast_next_4q(df):
    rev_growth = df["Total Revenue"].pct_change().mean()
    op_growth = df["Operating Income"].pct_change().mean()
    last_rev = df["Total Revenue"].iloc[-1]
    last_op = df["Operating Income"].iloc[-1]
    forecasts=[]
    for i in range(4):
        last_rev *= (1 + rev_growth)
        last_op *= (1 + op_growth)
        forecasts.append([last_rev, last_op])
    return pd.DataFrame(forecasts, columns=["Revenue Forecast", "Operating Income Forecast"], index=[f"Q+{i+1}" for i in range(4)])

# ==============================
# PER, RS 점수
# ==============================
def get_per(ticker):
    stock = yf.Ticker(ticker)
    return stock.fast_info.get("trailingPE", None)

def get_rs_score(ticker, spy_return):
    stock = yf.Ticker(ticker)
    hist = stock.history(period="1y")["Close"]
    stock_return = (hist.iloc[-1] - hist.iloc[0]) / hist.iloc[0]
    if spy_return == 0:
        return None
    return round((stock_return / spy_return) * 100, 2)

# ==============================
# 그래프 생성 (BytesIO)
# ==============================
def save_plot_annual(df, ticker):
    buf = BytesIO()
    ax = df.plot(kind="bar", figsize=(10,5), title=f"{ticker} 최근 3년 매출 및 영업이익")
    ax.set_ylabel("USD")
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(buf, format='png')
    plt.close()
    buf.seek(0)
    return buf

def save_plot_forecast(df, ticker):
    buf = BytesIO()
    ax = df.plot(kind="bar", figsize=(10,5), title=f"{ticker} 향후 4분기 실적 추정")
    ax.set_ylabel("USD")
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(buf, format='png')
    plt.close()
    buf.seek(0)
    return buf

# ==============================
# 추천 종목 함수
# ==============================
def get_sector_peers(ticker):
    stock = yf.Ticker(ticker)
    sector = stock.info.get("sector", None)
    if not sector:
        return []
    sector_map = {
        "Technology": ["MSFT","NVDA","GOOGL","META","AVGO","AMD","CRM","ORCL"],
        "Healthcare": ["LLY","JNJ","PFE","MRK","ABBV","TMO"],
        "Financial Services": ["JPM","BAC","GS","MS","BLK"],
        "Consumer Cyclical": ["AMZN","TSLA","HD","NKE","MCD"],
        "Communication Services": ["META","GOOGL","NFLX","DIS"],
        "Energy": ["XOM","CVX","COP","SLB"],
        "Industrials": ["BA","CAT","GE","RTX"]
    }
    return sector_map.get(sector, [])

def score_stock(ticker, spy_return):
    try:
        stock = yf.Ticker(ticker)
        info = stock.info
        per = info.get("trailingPE", None)
        if not per or per <= 0:
            return None
        hist = stock.history(period="1y")["Close"]
        stock_return = (hist.iloc[-1] - hist.iloc[0]) / hist.iloc[0]
        rs_score = (stock_return / spy_return) * 100
        quarterly = stock.quarterly_financials.loc["Total Revenue"]
        revenue_growth = quarterly.pct_change().mean()
        score = rs_score + (revenue_growth * 100)
        return {
            "ticker": ticker,
            "PER": round(per,2),
            "RS": round(rs_score,2),
            "RevenueGrowth": round(revenue_growth*100,2),
            "Score": round(score,2)
        }
    except:
        return None

def recommend_stocks(base_ticker):
    spy = yf.Ticker("SPY")
    spy_hist = spy.history(period="1y")["Close"]
    spy_return = (spy_hist.iloc[-1] - spy_hist.iloc[0]) / spy_hist.iloc[0]
    peers = get_sector_peers(base_ticker)
    scored = []
    for t in peers:
        result = score_stock(t, spy_return)
        if result:
            scored.append(result)
    if not scored:
        return pd.DataFrame()
    df = pd.DataFrame(scored)
    df = df.sort_values("Score", ascending=False)
    return df.head(3)

# ==============================
# GPT 보고서 함수
# ==============================
def generate_report(ticker, annual_df, quarterly_df, forecast_df, per, rs_score):
    report_prompt = f"""
    다음은 {ticker}의 실적/지표입니다.

    [최근 3개년 연간 실적]
    {annual_df}

    [최근 4개 분기 실적]
    {quarterly_df}

    [향후 4분기 추정치]
    {forecast_df}

    [PER]
    {per}

    [RS 점수 (시장 대비 상대강도)]
    {rs_score}

    위 정보를 기반으로:
    1) 최근 매출/영업이익 추이 분석
    2) 향후 실적 전망 및 이유
    3) PER 해석
    4) RS 점수 해석
    5) 리스크 요인
    6) 투자포인트 3가지

    한국어 리서치 보고서 형식으로 작성하세요.
    """
    return llm.invoke(report_prompt).content

def generate_recommendation_report(base_ticker, recommended_df):
    rec_prompt = f"""
    기준 종목: {base_ticker}

    추천 후보:
    {recommended_df}

    위 데이터를 기반으로:
    1) 왜 이 종목들이 더 매력적인지
    2) 성장성 비교
    3) 밸류에이션 비교
    4) 투자 시 유의점

    한국어로 간단한 투자 의견 작성
    """
    return llm.invoke(rec_prompt).content

# ==============================
# DOCX 저장
# ==============================
def save_report_docx(ticker, report_text, rec_report_text, annual_img, forecast_img):
    download_path = str(Path.home() / "Downloads")
    file_path = os.path.join(download_path, f"{ticker}_분석보고서.docx")
    doc = Document()
    doc.add_heading(f"{ticker} 분석 보고서", 0)
    doc.add_paragraph(report_text)
    doc.add_picture(annual_img, width=Inches(6))
    doc.add_picture(forecast_img, width=Inches(6))
    doc.add_paragraph("\n추천 종목 분석")
    doc.add_paragraph(rec_report_text)
    doc.save(file_path)
    return file_path

# ==============================
# 메인 분석 함수 (멀티 티커)
# ==============================
def analyze_stock(tickers):
    tickers = [t.strip().upper() for t in tickers.split(",")]
    for ticker in tickers:
        print(f"\n📊 {ticker} 분석 시작...")
        try:
            annual_df = get_annual_financials(ticker)
            quarterly_df = get_quarterly_actuals(ticker)
            forecast_df = forecast_next_4q(quarterly_df)
            per = get_per(ticker)
            spy = yf.Ticker("SPY")
            spy_hist = spy.history(period="1y")["Close"]
            spy_return = (spy_hist.iloc[-1] - spy_hist.iloc[0]) / spy_hist.iloc[0]
            rs_score = get_rs_score(ticker, spy_return)
        except Exception as e:
            print("데이터 수집 실패:", e)
            continue
        
        # 그래프 이미지 생성 (BytesIO)
        annual_img = save_plot_annual(annual_df, ticker)
        forecast_img = save_plot_forecast(forecast_df, ticker)
        
        report_text = generate_report(ticker, annual_df, quarterly_df, forecast_df, per, rs_score)
        recommended_df = recommend_stocks(ticker)
        rec_report_text = generate_recommendation_report(ticker, recommended_df) if not recommended_df.empty else "추천 종목이 없습니다."
        
        # DOCX 저장
        docx_file = save_report_docx(ticker, report_text, rec_report_text, annual_img, forecast_img)
        print(f"\n✅ {ticker} DOCX 보고서 생성 완료!")
        print(f"DOCX 파일 경로: {docx_file}")

# ==============================
# 사용자 입력 실행
# ==============================
if __name__ == "__main__":
    tickers_input = input("분석할 티커를 입력하세요 (예: AAPL, MSFT, NVDA): ")
    analyze_stock(tickers_input)